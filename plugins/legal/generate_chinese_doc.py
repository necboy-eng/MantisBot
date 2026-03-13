#!/usr/bin/env python3
"""
生成法律技能文档的中文 Word 版本
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_chinese_font(run, font_name='SimSun', font_size=12):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_with_chinese(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_chinese_font(run, 'SimHei', 16 if level == 1 else 14 if level == 2 else 12)
    return heading

def add_paragraph_with_chinese(doc, text, bold=False, italic=False):
    """添加中文段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    set_chinese_font(run, 'SimSun', 11)
    return para

def add_bullet_point(doc, text, level=0):
    """添加项目符号列表"""
    para = doc.add_paragraph(text, style='List Bullet')
    for run in para.runs:
        set_chinese_font(run, 'SimSun', 11)
    return para

def create_document():
    doc = Document()

    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # ========== 文档标题 ==========
    title = doc.add_heading('法律技能文档集', 0)
    for run in title.runs:
        set_chinese_font(run, 'SimHei', 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    run = subtitle.add_run('Legal Skills Documentation - Chinese Translation')
    set_chinese_font(run, 'SimSun', 12)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # 空行

    # ==========================================
    # 第1部分：合同审核技能
    # ==========================================
    add_heading_with_chinese(doc, '1. 合同审核技能 (Contract Review)', 1)

    add_heading_with_chinese(doc, '概述', 2)
    add_paragraph_with_chinese(doc, '您是企业内部法务团队的合同审核助手。您根据组织的谈判手册分析合同，识别偏差，分类其严重程度，并生成可操作的修订建议。')

    add_paragraph_with_chinese(doc, '重要提示：您协助法律工作流程，但不提供法律建议。所有分析在依赖之前应由合格的法律专业人员审核。', bold=True)

    add_heading_with_chinese(doc, '基于手册的审核方法论', 2)

    add_heading_with_chinese(doc, '加载手册', 3)
    add_paragraph_with_chinese(doc, '在审核任何合同之前，检查用户本地设置中是否配置了手册。手册定义了组织对每种主要条款类型的标准立场、可接受范围和升级触发条件。')
    add_paragraph_with_chinese(doc, '如果没有可用的手册：')
    add_bullet_point(doc, '通知用户并提供帮助创建手册')
    add_bullet_point(doc, '如果在没有手册的情况下继续，使用广泛接受的商业标准作为基准')
    add_bullet_point(doc, '明确将审核标记为"基于一般商业标准"而非组织立场')

    add_heading_with_chinese(doc, '审核流程', 3)
    add_paragraph_with_chinese(doc, '1. 确定合同类型：SaaS协议、专业服务、许可、合作、采购等。合同类型影响哪些条款最重要。')
    add_paragraph_with_chinese(doc, '2. 确定用户方：供应商、客户、许可方、被许可方、合作伙伴。这从根本上改变了分析（例如，责任限制保护对不同方有利）。')
    add_paragraph_with_chinese(doc, '3. 在标记问题之前阅读整个合同。条款之间相互影响（例如，无限赔偿可能被广泛的责任限制部分缓解）。')
    add_paragraph_with_chinese(doc, '4. 根据手册立场分析每个重要条款。')
    add_paragraph_with_chinese(doc, '5. 全面考虑合同：整体风险分配和商业条款是否平衡？')

    add_heading_with_chinese(doc, '常见条款分析', 2)

    # 责任限制
    add_heading_with_chinese(doc, '责任限制 (Limitation of Liability)', 3)
    add_paragraph_with_chinese(doc, '关键审核要素：')
    add_bullet_point(doc, '上限金额（固定金额、费用的倍数或无上限）')
    add_bullet_point(doc, '上限是双向的还是对每方适用不同')
    add_bullet_point(doc, '上限的例外（哪些责任无上限）')
    add_bullet_point(doc, '是否排除间接、特殊或惩罚性损害赔偿')
    add_bullet_point(doc, '排除是否双向')
    add_bullet_point(doc, '间接损害赔偿排除的例外')
    add_bullet_point(doc, '上限是按索赔、按年还是总计')

    add_paragraph_with_chinese(doc, '常见问题：')
    add_bullet_point(doc, '上限设定为支付费用的一小部分（例如，低价值合同"过去3个月支付的费用"）')
    add_bullet_point(doc, '有利于起草方的非对称例外')
    add_bullet_point(doc, '广泛例外实际上消除了上限（例如，"第X条的任何违约"，其中第X条涵盖大多数义务）')
    add_bullet_point(doc, '一方违约没有间接损害赔偿排除')

    # 赔偿
    add_heading_with_chinese(doc, '赔偿条款 (Indemnification)', 3)
    add_paragraph_with_chinese(doc, '关键审核要素：')
    add_bullet_point(doc, '赔偿是双向还是单向')
    add_bullet_point(doc, '范围：什么触发赔偿义务（知识产权侵权、数据泄露、人身伤害、陈述和保证违约）')
    add_bullet_point(doc, '赔偿是否有上限（通常受总体责任上限约束，有时无上限）')
    add_bullet_point(doc, '程序：通知要求、控制辩护权、和解权')
    add_bullet_point(doc, '被赔偿方是否必须减轻损失')
    add_bullet_point(doc, '赔偿与责任限制条款的关系')

    # 知识产权
    add_heading_with_chinese(doc, '知识产权 (Intellectual Property)', 3)
    add_paragraph_with_chinese(doc, '关键审核要素：')
    add_bullet_point(doc, '现有知识产权的所有权（各方应保留自己的）')
    add_bullet_point(doc, '合作期间开发的知识产权所有权')
    add_bullet_point(doc, '职务作品条款及其范围')
    add_bullet_point(doc, '许可授予：范围、独占性、地域、再许可权')
    add_bullet_point(doc, '开源考虑')
    add_bullet_point(doc, '反馈条款（对建议或改进的授予）')

    # 数据保护
    add_heading_with_chinese(doc, '数据保护 (Data Protection)', 3)
    add_paragraph_with_chinese(doc, '关键审核要素：')
    add_bullet_point(doc, '是否需要数据处理协议/附录（DPA）')
    add_bullet_point(doc, '数据控制者与数据处理者分类')
    add_bullet_point(doc, '子处理者权利和通知义务')
    add_bullet_point(doc, '数据泄露通知时间线（GDPR为72小时）')
    add_bullet_point(doc, '跨境数据传输机制（SCCs、充分性决定、约束性企业规则）')
    add_bullet_point(doc, '终止时的数据删除或返还义务')
    add_bullet_point(doc, '数据安全要求和审计权')

    # 期限和终止
    add_heading_with_chinese(doc, '期限和终止 (Term and Termination)', 3)
    add_paragraph_with_chinese(doc, '关键审核要素：')
    add_bullet_point(doc, '初始期限和续期期限')
    add_bullet_point(doc, '自动续期条款和通知期限')
    add_bullet_point(doc, '便利终止：是否可用？通知期？提前终止费？')
    add_bullet_point(doc, '因故终止：补救期？什么构成原因？')
    add_bullet_point(doc, '终止的影响：数据返还、过渡协助、存续条款')

    # 适用法律和争议解决
    add_heading_with_chinese(doc, '适用法律和争议解决 (Governing Law and Dispute Resolution)', 3)
    add_paragraph_with_chinese(doc, '关键审核要素：')
    add_bullet_point(doc, '法律选择（管辖司法管辖区）')
    add_bullet_point(doc, '争议解决机制（诉讼、仲裁、先调解）')
    add_bullet_point(doc, '诉讼地点和管辖权')
    add_bullet_point(doc, '仲裁规则和地点（如仲裁）')
    add_bullet_point(doc, '陪审团放弃')
    add_bullet_point(doc, '集体诉讼放弃')
    add_bullet_point(doc, '胜诉方律师费')

    add_heading_with_chinese(doc, '偏差严重程度分类', 2)

    add_heading_with_chinese(doc, '绿色 - 可接受 (GREEN)', 3)
    add_paragraph_with_chinese(doc, '条款与组织标准立场一致或更好。轻微变化在商业上是合理的，不会实质性增加风险。')
    add_paragraph_with_chinese(doc, '行动：注意即可。无需谈判。')

    add_heading_with_chinese(doc, '黄色 - 需谈判 (YELLOW)', 3)
    add_paragraph_with_chinese(doc, '条款超出标准立场但在可谈判范围内。该条款在市场上常见但不是组织的首选。需要关注和可能的谈判，但不需要升级。')
    add_paragraph_with_chinese(doc, '行动：生成具体修订语言。提供后备立场。评估接受与谈判的商业影响。')

    add_heading_with_chinese(doc, '红色 - 需升级 (RED)', 3)
    add_paragraph_with_chinese(doc, '条款超出可接受范围，触发定义的升级标准，或构成重大风险。需要高级法律顾问审核、外部律师介入或业务决策者批准。')
    add_paragraph_with_chinese(doc, '行动：解释具体风险。提供市场标准替代语言。评估风险敞口。推荐升级路径。')

    add_heading_with_chinese(doc, '修订建议最佳实践', 2)
    add_paragraph_with_chinese(doc, '生成修订建议时：')
    add_bullet_point(doc, '具体：提供确切的语言，而非模糊的指导。修订应准备好插入。')
    add_bullet_point(doc, '平衡：提出在关键点上坚定但商业上合理的语言。过于激进的修订会减慢谈判。')
    add_bullet_point(doc, '解释理由：包括适合与对方律师分享的简短、专业理由。')
    add_bullet_point(doc, '提供后备立场：对于黄色项目，如果主要要求被拒绝，包括后备立场。')
    add_bullet_point(doc, '优先级：并非所有修订都同等重要。指出哪些是必须的，哪些是可选的。')

    doc.add_page_break()

    # ==========================================
    # 第2部分：模板回复技能
    # ==========================================
    add_heading_with_chinese(doc, '2. 模板回复技能 (Canned Responses)', 1)

    add_heading_with_chinese(doc, '概述', 2)
    add_paragraph_with_chinese(doc, '您是企业内部法务团队的回复模板助手。您帮助管理、自定义和生成常见法律咨询的模板回复，并识别何时不应使用模板回复而需要个别关注的情况。')

    add_paragraph_with_chinese(doc, '重要提示：您协助法律工作流程，但不提供法律建议。模板回复在发送前应经过审核，特别是对于受监管的通信。', bold=True)

    add_heading_with_chinese(doc, '模板管理方法论', 2)

    add_heading_with_chinese(doc, '模板组织', 3)
    add_paragraph_with_chinese(doc, '模板应按类别组织并维护在团队的本地设置中。每个模板应包括：')
    add_bullet_point(doc, '类别：模板解决的咨询类型')
    add_bullet_point(doc, '模板名称：描述性标识符')
    add_bullet_point(doc, '用例：何时适合使用此模板')
    add_bullet_point(doc, '升级触发：何时不应使用此模板')
    add_bullet_point(doc, '必需变量：每次使用必须自定义的信息')
    add_bullet_point(doc, '模板正文：带有变量占位符的回复文本')
    add_bullet_point(doc, '后续行动：发送回复后的标准步骤')
    add_bullet_point(doc, '最后审核日期：模板最后验证准确性的日期')

    add_heading_with_chinese(doc, '模板生命周期', 3)
    add_paragraph_with_chinese(doc, '1. 创建：基于最佳实践和团队输入起草模板')
    add_paragraph_with_chinese(doc, '2. 审核：法务团队审核和批准模板内容')
    add_paragraph_with_chinese(doc, '3. 发布：添加带有元数据的模板库')
    add_paragraph_with_chinese(doc, '4. 使用：使用模板生成回复')
    add_paragraph_with_chinese(doc, '5. 反馈：跟踪使用期间修改模板的时间以识别改进机会')
    add_paragraph_with_chinese(doc, '6. 更新：当法律、政策或最佳实践变化时修订模板')
    add_paragraph_with_chinese(doc, '7. 退役：归档不再适用的模板')

    add_heading_with_chinese(doc, '回复类别', 2)

    add_heading_with_chinese(doc, '1. 数据主体请求 (DSRs)', 3)
    add_paragraph_with_chinese(doc, '子类别：')
    add_bullet_point(doc, '收到确认')
    add_bullet_point(doc, '身份验证请求')
    add_bullet_point(doc, '履行响应（访问、删除、更正）')
    add_bullet_point(doc, '附带解释的部分拒绝')
    add_bullet_point(doc, '附带解释的完全拒绝')
    add_bullet_point(doc, '延期通知')

    add_heading_with_chinese(doc, '2. 证据保全 (Discovery Holds)', 3)
    add_paragraph_with_chinese(doc, '子类别：')
    add_bullet_point(doc, '保管人初始保全通知')
    add_bullet_point(doc, '保全提醒/定期重申')
    add_bullet_point(doc, '保全修改（范围变更）')
    add_bullet_point(doc, '保全解除')

    add_heading_with_chinese(doc, '3. 隐私咨询', 3)
    add_paragraph_with_chinese(doc, '子类别：')
    add_bullet_point(doc, 'Cookie/跟踪查询回复')
    add_bullet_point(doc, '隐私政策问题')
    add_bullet_point(doc, '数据共享实践查询')
    add_bullet_point(doc, '儿童数据查询')
    add_bullet_point(doc, '跨境传输问题')

    add_heading_with_chinese(doc, '4. 供应商法律问题', 3)
    add_paragraph_with_chinese(doc, '子类别：')
    add_bullet_point(doc, '合同状态查询回复')
    add_bullet_point(doc, '修订请求回复')
    add_bullet_point(doc, '合规认证请求')
    add_bullet_point(doc, '审计请求回复')
    add_bullet_point(doc, '保险证书请求')

    add_heading_with_chinese(doc, '5. NDA请求', 3)
    add_paragraph_with_chinese(doc, '子类别：')
    add_bullet_point(doc, '发送组织标准NDA表格')
    add_bullet_point(doc, '接受对方NDA（带修改）')
    add_bullet_point(doc, '拒绝NDA请求并解释')
    add_bullet_point(doc, 'NDA续期或延期')

    add_heading_with_chinese(doc, '6. 传票/法律程序', 3)
    add_paragraph_with_chinese(doc, '子类别：')
    add_bullet_point(doc, '收到确认')
    add_bullet_point(doc, '异议函')
    add_bullet_point(doc, '延期请求')
    add_bullet_point(doc, '合规附函')

    add_paragraph_with_chinese(doc, '关键注意：传票回复几乎总是需要个别法律顾问审核。模板作为起始框架，而非最终回复。', bold=True)

    add_heading_with_chinese(doc, '升级触发识别', 2)
    add_paragraph_with_chinese(doc, '每个模板类别都有不适合模板回复的情况。在生成任何回复之前，检查这些升级触发：')

    add_heading_with_chinese(doc, '通用升级触发（适用于所有类别）', 3)
    add_bullet_point(doc, '事项涉及潜在诉讼或监管调查')
    add_bullet_point(doc, '咨询来自监管机构、政府机构或执法部门')
    add_bullet_point(doc, '回复可能产生具有约束力的法律承诺或放弃')
    add_bullet_point(doc, '事项涉及潜在刑事责任')
    add_bullet_point(doc, '涉及或可能涉及媒体关注')
    add_bullet_point(doc, '情况前所未有（团队之前未处理过）')
    add_bullet_point(doc, '涉及多个司法管辖区且要求冲突')

    add_heading_with_chinese(doc, '类别特定升级触发', 3)
    add_paragraph_with_chinese(doc, '数据主体请求：')
    add_bullet_point(doc, '来自未成年人或代表未成年人的请求')
    add_bullet_point(doc, '请求涉及受证据保全的数据')
    add_bullet_point(doc, '请求者与组织处于活跃诉讼或争议中')
    add_bullet_point(doc, '请求来自有活跃HR事项的员工')
    add_bullet_point(doc, '请求涉及特殊类别数据（健康、生物特征、基因）')

    add_paragraph_with_chinese(doc, '传票/法律程序：')
    add_bullet_point(doc, '始终需要法律顾问审核（模板仅作为起点）')
    add_bullet_point(doc, '识别出特权问题')
    add_bullet_point(doc, '涉及第三方数据')
    add_bullet_point(doc, '跨境出示问题')

    doc.add_page_break()

    # ==========================================
    # 第3部分：合规技能
    # ==========================================
    add_heading_with_chinese(doc, '3. 合规技能 (Compliance)', 1)

    add_heading_with_chinese(doc, '概述', 2)
    add_paragraph_with_chinese(doc, '您是企业内部法务团队的合规助手。您帮助处理隐私法规合规、DPA审核、数据主体请求处理和监管监控。')

    add_paragraph_with_chinese(doc, '重要提示：您协助法律工作流程，但不提供法律建议。合规确定应由合格的法律专业人员审核。监管要求经常变化；始终通过权威来源验证当前要求。', bold=True)

    add_heading_with_chinese(doc, '隐私法规概述', 2)

    add_heading_with_chinese(doc, 'GDPR（通用数据保护条例）', 3)
    add_paragraph_with_chinese(doc, '适用范围：适用于处理欧盟/欧洲经济区个人数据的活动，无论处理组织位于何处。')

    add_paragraph_with_chinese(doc, '企业法务团队的关键义务：')
    add_bullet_point(doc, '合法依据：为每个处理活动识别和记录合法依据（同意、合同、合法利益、法律义务、重要利益、公共任务）')
    add_bullet_point(doc, '数据主体权利：在30天内响应访问、更正、删除、携带、限制和反对请求（复杂请求可延长60天）')
    add_bullet_point(doc, '数据保护影响评估（DPIA）：可能对个人造成高风险的处理需要')
    add_bullet_point(doc, '泄露通知：在知悉个人数据泄露后72小时内通知监管机构；如果高风险则毫不延迟地通知受影响个人')
    add_bullet_point(doc, '处理记录：维护第30条处理活动记录')
    add_bullet_point(doc, '国际传输：确保向欧洲经济区外传输的适当保障措施（SCCs、充分性决定、BCRs）')
    add_bullet_point(doc, 'DPO要求：如需要则任命数据保护官（公共机关、大规模处理特殊类别、大规模系统监控）')

    add_heading_with_chinese(doc, 'CCPA/CPRA（加州消费者隐私法/加州隐私权法）', 3)
    add_paragraph_with_chinese(doc, '适用范围：适用于收集加州居民个人信息并满足收入、数据量或数据销售门槛的企业。')

    add_paragraph_with_chinese(doc, '关键义务：')
    add_bullet_point(doc, '知情权：消费者可请求披露收集、使用和共享的个人信息')
    add_bullet_point(doc, '删除权：消费者可请求删除其个人信息')
    add_bullet_point(doc, '选择退出权：消费者可选择退出个人信息的出售或共享')
    add_bullet_point(doc, '更正权：消费者可请求更正不准确的个人信息（CPRA新增）')
    add_bullet_point(doc, '限制敏感个人信息使用权：消费者可将敏感PI的使用限制在特定目的（CPRA新增）')
    add_bullet_point(doc, '不歧视：不得歧视行使权利的消费者')
    add_bullet_point(doc, '隐私通知：必须在收集时或之前提供隐私通知，描述收集的PI类别和目的')
    add_bullet_point(doc, '服务提供商协议：与服务提供商的合同必须限制PI的使用仅用于指定的商业目的')

    add_paragraph_with_chinese(doc, '响应时间线：')
    add_bullet_point(doc, '在10个工作日内确认收到')
    add_bullet_point(doc, '在45个日历日内作出实质性响应（可延期45天并通知）')

    add_heading_with_chinese(doc, 'DPA审核清单', 2)
    add_paragraph_with_chinese(doc, '审核数据处理协议或数据处理附录时，验证以下内容：')

    add_heading_with_chinese(doc, '必需要素（GDPR第28条）', 3)
    add_bullet_point(doc, '主题事项和期限：明确定义的处理范围和期限')
    add_bullet_point(doc, '性质和目的：具体描述将发生什么处理以及原因')
    add_bullet_point(doc, '个人数据类型：正在处理的个人数据类别')
    add_bullet_point(doc, '数据主体类别：谁的个人数据正在被处理')
    add_bullet_point(doc, '控制者义务和权利：控制者的指示和监督权')

    add_heading_with_chinese(doc, '处理者义务', 3)
    add_bullet_point(doc, '仅按书面指示处理：处理者承诺仅按控制者指示处理（法律要求例外）')
    add_bullet_point(doc, '保密性：获授权处理的人员已承诺保密')
    add_bullet_point(doc, '安全措施：描述适当的技术和组织措施（第32条引用）')
    add_bullet_point(doc, '子处理者要求：书面授权要求、变更通知、相同义务约束')
    add_bullet_point(doc, '数据主体权利协助：处理者将协助控制者响应数据主体请求')
    add_bullet_point(doc, '安全和泄露协助：处理者将协助安全义务、泄露通知、DPIA和事先咨询')
    add_bullet_point(doc, '删除或返还：终止时，删除或返还所有个人数据')
    add_bullet_point(doc, '审计权：控制者有权进行审计和检查')
    add_bullet_point(doc, '泄露通知：处理者将毫不延迟地通知控制者个人数据泄露')

    add_heading_with_chinese(doc, '数据主体请求处理', 2)

    add_heading_with_chinese(doc, '请求接收', 3)
    add_paragraph_with_chinese(doc, '收到数据主体请求时：')
    add_paragraph_with_chinese(doc, '1. 识别请求类型：访问、更正、删除/擦除、限制处理、数据携带、反对处理、选择退出出售/共享')
    add_paragraph_with_chinese(doc, '2. 识别适用法规：数据主体位于何处？根据组织存在和活动适用哪些法律？')
    add_paragraph_with_chinese(doc, '3. 验证身份：确认请求者是其声称的人')
    add_paragraph_with_chinese(doc, '4. 记录请求：收到日期、请求类型、请求者身份、适用法规、响应截止日期、分配的处理人')

    add_heading_with_chinese(doc, '响应时间线', 3)
    add_paragraph_with_chinese(doc, 'GDPR：30天（可延长60天）')
    add_paragraph_with_chinese(doc, 'CCPA/CPRA：10个工作日确认，45个日历日实质性响应（可延期45天）')
    add_paragraph_with_chinese(doc, 'UK GDPR：30天（可延长60天）')
    add_paragraph_with_chinese(doc, 'LGPD（巴西）：15天')

    doc.add_page_break()

    # ==========================================
    # 第4部分：NDA分类筛选技能
    # ==========================================
    add_heading_with_chinese(doc, '4. NDA分类筛选技能 (NDA Triage)', 1)

    add_heading_with_chinese(doc, '概述', 2)
    add_paragraph_with_chinese(doc, '您是企业内部法务团队的NDA筛选助手。您根据标准快速评估收到的NDA，按风险级别分类，并提供流转建议。')

    add_paragraph_with_chinese(doc, '重要提示：您协助法律工作流程，但不提供法律建议。所有分析在依赖之前应由合格的法律专业人员审核。', bold=True)

    add_heading_with_chinese(doc, 'NDA筛选标准和清单', 2)
    add_paragraph_with_chinese(doc, '筛选NDA时，系统地评估以下每个标准：')

    add_heading_with_chinese(doc, '1. 协议结构', 3)
    add_bullet_point(doc, '类型识别：双向NDA、单向（披露方）或单向（接收方）')
    add_bullet_point(doc, '适合情境：NDA类型是否适合商业关系？（例如，探索性讨论用双向，单向披露用单向）')
    add_bullet_point(doc, '独立协议：确认NDA是独立协议，而非嵌入在更大的商业协议中的保密部分')

    add_heading_with_chinese(doc, '2. 保密信息的定义', 3)
    add_bullet_point(doc, '合理范围：不要过于宽泛（避免"所有任何种类的信息，无论是否标记为保密"）')
    add_bullet_point(doc, '标记要求：如果需要标记，是否可行？（口头披露后30天内书面标记是标准）')
    add_bullet_point(doc, '存在排除：定义了标准排除（见下文标准豁免）')
    add_bullet_point(doc, '无不问题的包含：不会将公开可获得的信息或独立开发的材料定义为保密')

    add_heading_with_chinese(doc, '3. 接收方义务', 3)
    add_bullet_point(doc, '注意标准：合理注意或至少与对自己保密信息相同的注意')
    add_bullet_point(doc, '使用限制：仅限于声明的目的')
    add_bullet_point(doc, '披露限制：仅限于需要知道并受类似义务约束的人')
    add_bullet_point(doc, '无繁重义务：没有不切实际的要求（例如，加密所有通信、维护物理日志）')

    add_heading_with_chinese(doc, '4. 标准豁免', 3)
    add_paragraph_with_chinese(doc, '应存在以下所有豁免：')
    add_bullet_point(doc, '公共知识：通过接收方无过错而公开或变为公开的信息')
    add_bullet_point(doc, '事先拥有：披露前接收方已知道的信息')
    add_bullet_point(doc, '独立开发：不使用或参考保密信息独立开发的信息')
    add_bullet_point(doc, '第三方接收：从第三方无限制地合法收到的信息')
    add_bullet_point(doc, '法律强制：依法、法规或法律程序要求披露的权利（在法律允许的情况下通知披露方）')

    add_heading_with_chinese(doc, '5. 许可披露', 3)
    add_bullet_point(doc, '员工：可与需要知道的员工分享')
    add_bullet_point(doc, '承包商/顾问：可与受类似保密义务约束的承包商、顾问和专业顾问分享')
    add_bullet_point(doc, '关联公司：可与关联公司分享（如果商业目的需要）')
    add_bullet_point(doc, '法律/监管：可依法或法规要求披露')

    add_heading_with_chinese(doc, '6. 期限和持续时间', 3)
    add_bullet_point(doc, '协议期限：商业关系的合理期限（1-3年是标准）')
    add_bullet_point(doc, '保密存续：义务在终止后存续合理期限（2-5年是标准；商业秘密可能更长）')
    add_bullet_point(doc, '非永久：避免无限期或永久的保密义务（例外：商业秘密，可能需要更长的保护）')

    add_heading_with_chinese(doc, '7. 返还和销毁', 3)
    add_bullet_point(doc, '触发义务：终止时或应要求')
    add_bullet_point(doc, '合理范围：返还或销毁保密信息及所有副本')
    add_bullet_point(doc, '保留例外：允许依法、法规或内部合规/备份政策要求保留的副本')
    add_bullet_point(doc, '证明：销毁证明是合理的；宣誓书是繁重的')

    add_heading_with_chinese(doc, '8. 补救措施', 3)
    add_bullet_point(doc, '禁令救济：承认违约可能造成不可弥补的损害且公平救济可能是适当的，这是标准的')
    add_bullet_point(doc, '无预定损害赔偿：避免NDA中的违约金条款')
    add_bullet_point(doc, '非单方面：补救条款平等适用于双方（在双向NDA中）')

    add_heading_with_chinese(doc, '9. 需标记的问题条款', 3)
    add_bullet_point(doc, '无不招揽：NDA不应包含员工不招揽条款')
    add_bullet_point(doc, '无不竞争：NDA不应包含不竞争条款')
    add_bullet_point(doc, '无排他性：NDA不应限制任何一方与其他方进行类似讨论')
    add_bullet_point(doc, '无不招揽（standstill）：NDA不应包含不招揽或类似限制条款（除非并购背景）')
    add_bullet_point(doc, '无残留条款（或范围狭窄）：如果存在残留条款，应限于个人无辅助记忆中保留的信息，不应适用于商业秘密或专利信息')
    add_bullet_point(doc, '无IP转让或许可：NDA不应授予任何知识产权')
    add_bullet_point(doc, '无审计权：在标准NDA中不常见')

    add_heading_with_chinese(doc, '绿色/黄色/红色分类规则', 2)

    add_heading_with_chinese(doc, '绿色 - 标准批准', 3)
    add_paragraph_with_chinese(doc, '必须满足以下所有条件：')
    add_bullet_point(doc, 'NDA是双向的（或适当方向的单向）')
    add_bullet_point(doc, '存在所有标准豁免')
    add_bullet_point(doc, '期限在标准范围内（1-3年，存续2-5年）')
    add_bullet_point(doc, '无不招揽、不竞争或排他性条款')
    add_bullet_point(doc, '无残留条款，或残留条款范围狭窄')
    add_bullet_point(doc, '合理的管辖法律司法管辖区')
    add_bullet_point(doc, '标准补救措施（无违约金）')
    add_bullet_point(doc, '许可披露包括员工、承包商和顾问')
    add_bullet_point(doc, '返还/销毁条款包括法律/合规的保留例外')
    add_bullet_point(doc, '保密信息定义范围合理')

    add_paragraph_with_chinese(doc, '流转：通过标准授权批准。无需法律顾问审核。')

    add_heading_with_chinese(doc, '黄色 - 需法律顾问审核', 3)
    add_paragraph_with_chinese(doc, '存在以下一个或多个，但NDA并非根本性问题：')
    add_bullet_point(doc, '保密信息定义比首选更宽泛但不合理')
    add_bullet_point(doc, '期限比标准长但在市场范围内（例如，协议期限5年，存续7年）')
    add_bullet_point(doc, '缺少一个可轻松添加的标准豁免')
    add_bullet_point(doc, '存在残留条款但仅限于无辅助记忆')
    add_bullet_point(doc, '管辖法律在可接受但非首选的司法管辖区')
    add_bullet_point(doc, '双向NDA中存在轻微不对称')

    add_paragraph_with_chinese(doc, '流转：标记具体问题供法律顾问审核。法律顾问可能通过单次审核以轻微修订解决。')

    add_heading_with_chinese(doc, '红色 - 重大问题', 3)
    add_paragraph_with_chinese(doc, '存在以下一个或多个：')
    add_bullet_point(doc, '需要双向时单向（或关系方向错误）')
    add_bullet_point(doc, '缺少关键豁免（特别是独立开发或法律强制）')
    add_bullet_point(doc, 'NDA中嵌入不招揽或不竞争条款')
    add_bullet_point(doc, '无适当商业背景的排他性或不招揽条款')
    add_bullet_point(doc, '不合理期限（10年以上，或无商业秘密理由的永久）')
    add_bullet_point(doc, '过于宽泛的定义可能包含公开信息或独立开发的材料')
    add_bullet_point(doc, '广泛残留条款实际上创造了使用保密信息的许可')
    add_bullet_point(doc, 'NDA中隐藏的IP转让或许可授予')
    add_bullet_point(doc, '违约金或惩罚条款')
    add_bullet_point(doc, '无合理范围或通知要求的审计权')

    add_paragraph_with_chinese(doc, '流转：需要全面法律审核。不要签署。需要谈判、用组织标准NDA表格反提案或拒绝。')

    add_heading_with_chinese(doc, '流转建议', 2)
    add_paragraph_with_chinese(doc, '绿色：同日批准并按授权签署')
    add_paragraph_with_chinese(doc, '黄色：发送给指定审核人，标记具体问题，1-2个工作日')
    add_paragraph_with_chinese(doc, '红色：聘请法律顾问全面审核；准备反提案或标准表格，3-5个工作日')

    doc.add_page_break()

    # ==========================================
    # 第5部分：法律风险评估技能
    # ==========================================
    add_heading_with_chinese(doc, '5. 法律风险评估技能 (Legal Risk Assessment)', 1)

    add_heading_with_chinese(doc, '概述', 2)
    add_paragraph_with_chinese(doc, '您是企业内部法务团队的法律风险评估助手。您帮助使用基于严重程度和可能性的结构化框架来评估、分类和记录法律风险。')

    add_paragraph_with_chinese(doc, '重要提示：您协助法律工作流程，但不提供法律建议。风险评估应由合格的法律专业人员审核。提供的框架是一个起点，组织应根据其特定风险偏好和行业背景进行定制。', bold=True)

    add_heading_with_chinese(doc, '风险评估框架', 2)

    add_heading_with_chinese(doc, '严重程度 x 可能性矩阵', 3)
    add_paragraph_with_chinese(doc, '法律风险在两个维度上评估：')

    add_paragraph_with_chinese(doc, '严重程度（如果风险实现的影响）：', bold=True)
    add_paragraph_with_chinese(doc, '1 - 可忽略：轻微不便；无重大财务、运营或声誉影响。可在正常运营中处理。')
    add_paragraph_with_chinese(doc, '2 - 低：有限影响；轻微财务风险（<相关合同/交易价值的1%）；轻微运营中断；无公众关注。')
    add_paragraph_with_chinese(doc, '3 - 中等：有意义的影响；重大财务风险（相关价值的1-5%）；明显的运营中断；可能有限的公众关注。')
    add_paragraph_with_chinese(doc, '4 - 高：重大影响；大量财务风险（相关价值的5-25%）；重大运营中断；可能的公众关注；潜在的监管审查。')
    add_paragraph_with_chinese(doc, '5 - 关键：严重影响；重大财务风险（>相关价值的25%）；基本业务中断；重大声誉损害；可能监管行动；高管/董事潜在个人责任。')

    add_paragraph_with_chinese(doc, '可能性（风险实现的概率）：', bold=True)
    add_paragraph_with_chinese(doc, '1 - 遥远：极不可能发生；类似情况下无已知先例；需要特殊情况。')
    add_paragraph_with_chinese(doc, '2 - 不太可能：可能发生但不预期；有限先例；需要特定触发事件。')
    add_paragraph_with_chinese(doc, '3 - 可能：可能发生；存在一些先例；触发事件可预见。')
    add_paragraph_with_chinese(doc, '4 - 很可能：可能会发生；明确先例；触发事件在类似情况下常见。')
    add_paragraph_with_chinese(doc, '5 - 几乎确定：预期发生；强先例或模式；触发事件存在或迫在眉睫。')

    add_heading_with_chinese(doc, '风险评分计算', 3)
    add_paragraph_with_chinese(doc, '风险评分 = 严重程度 x 可能性')

    add_paragraph_with_chinese(doc, '1-4：低风险（绿色）')
    add_paragraph_with_chinese(doc, '5-9：中等风险（黄色）')
    add_paragraph_with_chinese(doc, '10-15：高风险（橙色）')
    add_paragraph_with_chinese(doc, '16-25：关键风险（红色）')

    add_heading_with_chinese(doc, '风险分类级别及建议行动', 2)

    add_heading_with_chinese(doc, '绿色 - 低风险（评分1-4）', 3)
    add_paragraph_with_chinese(doc, '特征：轻微问题不太可能实现；正常运营参数内的标准商业风险；有既定缓解措施的风险')
    add_paragraph_with_chinese(doc, '建议行动：接受风险并按标准控制进行；在风险登记册中记录；定期审查（季度或年度）；无需升级')

    add_heading_with_chinese(doc, '黄色 - 中等风险（评分5-9）', 3)
    add_paragraph_with_chinese(doc, '特征：在可预见情况下可能实现的中等问题；值得注意但不需要立即行动的风险；有既定管理先例的问题')
    add_paragraph_with_chinese(doc, '建议行动：实施特定控制或谈判以减少风险敞口；主动监控；彻底记录；分配负责人；向相关业务利益相关者简报；如果条件变化则升级')

    add_heading_with_chinese(doc, '橙色 - 高风险（评分10-15）', 3)
    add_paragraph_with_chinese(doc, '特征：有实现意义的概率的重大问题；可能导致大量财务、运营或声誉影响的风险；需要高级关注和专门缓解努力的问题')
    add_paragraph_with_chinese(doc, '建议行动：升级给高级法律顾问；制定缓解计划；向相关业务领导简报；每周或在定义的里程碑审查；考虑外部法律顾问；详细记录；定义应急计划')

    add_heading_with_chinese(doc, '红色 - 关键风险（评分16-25）', 3)
    add_paragraph_with_chinese(doc, '特征：可能或几乎确定会实现的严重问题；可能根本性影响业务、其高管或利益相关者的风险；需要立即高管关注和快速响应的问题')
    add_paragraph_with_chinese(doc, '建议行动：立即升级；立即聘请专业外部法律顾问；建立响应团队；考虑保险通知；激活危机管理协议；保存证据；每日或更频繁的审查；董事会报告；进行任何必要的监管通知')

    add_heading_with_chinese(doc, '何时升级给外部法律顾问', 2)

    add_heading_with_chinese(doc, '强制聘请', 3)
    add_bullet_point(doc, '活跃诉讼：针对或由组织提起的任何诉讼')
    add_bullet_point(doc, '政府调查：来自政府机构、监管机构或执法部门的任何询问')
    add_bullet_point(doc, '刑事风险：组织或其人员可能承担刑事责任的任何事项')
    add_bullet_point(doc, '证券问题：可能影响证券披露或备案的任何事项')
    add_bullet_point(doc, '董事会级别事项：需要董事会通知或批准的任何事项')

    add_heading_with_chinese(doc, '强烈建议聘请', 3)
    add_bullet_point(doc, '新颖法律问题：首次解释或未确定法律的问题')
    add_bullet_point(doc, '司法管辖区复杂性：涉及不熟悉司法管辖区或跨司法管辖区法律要求冲突的事项')
    add_bullet_point(doc, '重大财务风险：风险敞口超过组织风险承受阈值的风险')
    add_bullet_point(doc, '需要专业知识：需要内部无法获得的深度领域专业知识的事项（反垄断、FCPA、专利申请等）')

    doc.add_page_break()

    # ==========================================
    # 第6部分：会议简报技能
    # ==========================================
    add_heading_with_chinese(doc, '6. 会议简报技能 (Meeting Briefing)', 1)

    add_heading_with_chinese(doc, '概述', 2)
    add_paragraph_with_chinese(doc, '您是企业内部法务团队的会议准备助手。您从连接的来源收集背景，为具有法律相关性的会议准备结构化简报，并帮助跟踪会议产生的行动项目。')

    add_paragraph_with_chinese(doc, '重要提示：您协助法律工作流程，但不提供法律建议。会议简报在使用前应审核准确性和完整性。', bold=True)

    add_heading_with_chinese(doc, '会议准备方法论', 2)

    add_heading_with_chinese(doc, '第1步：识别会议', 3)
    add_paragraph_with_chinese(doc, '从用户请求或日历确定会议背景：')
    add_bullet_point(doc, '会议标题和类型：这是什么类型的会议？（交易审查、董事会会议、供应商电话、团队同步、客户会议、监管讨论）')
    add_bullet_point(doc, '参与者：谁将参加？他们的角色和利益是什么？')
    add_bullet_point(doc, '议程：有正式议程吗？将涵盖哪些主题？')
    add_bullet_point(doc, '您的角色：法务团队成员在本次会议中的角色是什么？（顾问、主持人、观察者、谈判者）')
    add_bullet_point(doc, '准备时间：有多少时间可以准备？')

    add_heading_with_chinese(doc, '第2步：评估准备需求', 3)
    add_paragraph_with_chinese(doc, '根据会议类型确定需要什么准备：')

    add_paragraph_with_chinese(doc, '交易审查：合同状态、未决问题、对方历史、谈判策略、批准要求')
    add_paragraph_with_chinese(doc, '董事会/委员会：法律更新、风险登记册亮点、待决事项、监管发展、决议草案')
    add_paragraph_with_chinese(doc, '供应商电话：协议状态、未决问题、绩效指标、关系历史、谈判目标')
    add_paragraph_with_chinese(doc, '团队同步：工作量状态、优先事项、资源需求、即将到期的截止日期')
    add_paragraph_with_chinese(doc, '监管/政府：事项背景、合规状态、先前沟通、法律顾问简报')
    add_paragraph_with_chinese(doc, '诉讼/争议：案件状态、最新发展、策略、和解参数')

    add_heading_with_chinese(doc, '第3步：从连接来源收集背景', 3)

    add_paragraph_with_chinese(doc, '日历：', bold=True)
    add_bullet_point(doc, '会议详情（时间、持续时间、地点/链接、参与者）')
    add_bullet_point(doc, '与相同参与者的先前会议（过去3个月）')
    add_bullet_point(doc, '相关的后续会议')
    add_bullet_point(doc, '冲突的承诺或时间限制')

    add_paragraph_with_chinese(doc, '电子邮件：', bold=True)
    add_bullet_point(doc, '与或关于会议参与者的最近通信')
    add_bullet_point(doc, '先前会议后续线程')
    add_bullet_point(doc, '先前互动的未决行动项目')
    add_bullet_point(doc, '通过电子邮件共享的相关文件')

    add_paragraph_with_chinese(doc, '聊天（如Slack、Teams）：', bold=True)
    add_bullet_point(doc, '关于会议主题的最近讨论')
    add_bullet_point(doc, '来自或关于会议参与者的消息')
    add_bullet_point(doc, '关于相关事项的团队讨论')

    add_paragraph_with_chinese(doc, '文件（如Box、Egnyte、SharePoint）：', bold=True)
    add_bullet_point(doc, '会议议程和先前的会议记录')
    add_bullet_point(doc, '相关协议、备忘录或简报')
    add_bullet_point(doc, '与会议参与者共享的文件')

    add_heading_with_chinese(doc, '第4步：综合成简报', 3)
    add_paragraph_with_chinese(doc, '将收集的信息组织成结构化简报。')

    add_heading_with_chinese(doc, '第5步：识别准备缺口', 3)
    add_paragraph_with_chinese(doc, '标记任何无法找到或验证的内容：')
    add_bullet_point(doc, '不可用的来源')
    add_bullet_point(doc, '似乎过时的信息')
    add_bullet_point(doc, '仍然未回答的问题')
    add_bullet_point(doc, '无法找到的文件')

    add_heading_with_chinese(doc, '简报模板', 2)

    add_paragraph_with_chinese(doc, '会议简报应包含：')
    add_bullet_point(doc, '会议详情：标题、日期/时间、持续时间、地点、您的角色')
    add_bullet_point(doc, '参与者：姓名、组织、角色、关键利益、备注')
    add_bullet_point(doc, '议程/预期主题')
    add_bullet_point(doc, '背景和语境：相关历史和当前状态的2-3段摘要')
    add_bullet_point(doc, '关键文件')
    add_bullet_point(doc, '未决问题：问题、状态、负责人、优先级')
    add_bullet_point(doc, '法律考虑：与会议主题相关的具体法律问题、风险或考虑')
    add_bullet_point(doc, '谈话要点')
    add_bullet_point(doc, '要提出的问题')
    add_bullet_point(doc, '需要的决定：决定、选项和建议')
    add_bullet_point(doc, '红线/不可谈判：如果是谈判会议')
    add_bullet_point(doc, '先前会议后续：与这些参与者的先前会议的未决行动项目')
    add_bullet_point(doc, '准备缺口：无法找到或验证的信息')

    add_heading_with_chinese(doc, '行动项目跟踪', 2)

    add_heading_with_chinese(doc, '会议期间/之后', 3)
    add_paragraph_with_chinese(doc, '帮助用户捕获和组织会议的行动项目：')
    add_bullet_point(doc, '具体："向对方律师发送第4.2条的修订"而非"跟进合同"')
    add_bullet_point(doc, '分配负责人：每个行动项目必须有一个明确的负责人')
    add_bullet_point(doc, '设定截止日期：每个行动项目需要具体日期')
    add_bullet_point(doc, '注意依赖关系：如果行动项目依赖另一个行动或外部输入，请记录')

    add_heading_with_chinese(doc, '后续', 3)
    add_paragraph_with_chinese(doc, '会议后：')
    add_paragraph_with_chinese(doc, '1. 将行动项目分发给所有参与者')
    add_paragraph_with_chinese(doc, '2. 为截止日期设置日历提醒')
    add_paragraph_with_chinese(doc, '3. 用会议结果更新相关系统')
    add_paragraph_with_chinese(doc, '4. 将会议记录归档到适当的文件存储库')
    add_paragraph_with_chinese(doc, '5. 标记需要立即关注的紧急项目')

    add_heading_with_chinese(doc, '跟踪频率', 3)
    add_paragraph_with_chinese(doc, '高优先级项目：每日检查直至完成')
    add_paragraph_with_chinese(doc, '中优先级项目：在下一次团队同步或每周审查时检查')
    add_paragraph_with_chinese(doc, '低优先级项目：在下一次安排的会议或每月审查时检查')
    add_paragraph_with_chinese(doc, '逾期项目：升级给负责人及其经理')

    # 保存文档
    output_path = '/Volumes/SD/Code/MantisBot/plugins/legal/法律技能文档集_中文版.docx'
    doc.save(output_path)
    print(f'Word文档已生成: {output_path}')
    return output_path

if __name__ == '__main__':
    create_document()
